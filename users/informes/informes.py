from graficas_informe import graficas_informe
from docx import Document
from docx.shared import Inches
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
#pip install python-docx


class informe_semanal():
    def __init__(self, empresa):
        self.empresa=empresa
        self.document = Document()

    def portada(self):
        self.document.add_paragraph('\n\n\n\n\n\n\n')
        t=self.document.add_heading('Informe Semanal '+self.empresa, 0)
        
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style = t.style
        title_style.font.size = Pt(40)

        self.document.add_paragraph('\n\n')
        self.document.add_picture(os.path.dirname(__file__)+'/imagenes/logo_mcreif.png', width=Inches(4.5))
        last_paragraph = self.document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag1(self, ingresos, ventas, cuadro_texto, graficos):
        self.document.add_heading('Temas generales', 1)
        self.document.add_heading('Ingresos semanales', 2)
        self.document.add_paragraph(
            'Ingresos: '+str(ingresos)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Ventas: '+str(ventas)+' unidades', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag2(self, ingresos, ventas, cuadro_texto, graficos):
        self.document.add_heading('Ingresos mensuales', 2)
        self.document.add_paragraph(
            'Ingresos: '+str(ingresos)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Ventas: '+str(ventas)+' unidades', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag3(self, obj_sell_in, obj_sell_out, cuadro_texto, graficos):
        self.document.add_heading('Objetivo sell in y sell out', 2)
        self.document.add_paragraph(
            'Objetivo sell in: '+str(obj_sell_in)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Objetivo sell out: '+str(obj_sell_out)+' €', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag4(self, cuadro_texto, graficos, buy_box=None, exceso_inventario=None, reabastecer=None):
        self.document.add_heading('Alertas', 2)
        if buy_box != None:
            self.document.add_paragraph(
                'Buy box perdida: '+str(buy_box)+' %', style='List Bullet'
            )
        if exceso_inventario != None:
            self.document.add_paragraph(
                'Unidades en exceso del inventario: '+str(exceso_inventario)+' unidades', style='List Bullet'
            )
        if reabastecer != None:
            self.document.add_paragraph(
                'Unidades a reabastecer del inventario: '+str(reabastecer)+' unidades', style='List Bullet'
            )
        
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag5(self, buy_box_stock, buy_box_precio, cuadro_texto, graficos):
        self.document.add_heading('Posicionamiento', 1)
        self.document.add_heading('Buy box', 2)
        self.document.add_paragraph(
            'Buy box perdida por stock : '+str(buy_box_stock)+' %', style='List Bullet'
        )
        self.document.add_paragraph(
            'Buy box perdida por precio: '+str(buy_box_precio)+' %', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag6(self, producto_mas_ingresos, producto_mas_vendido, cuadro_texto, graficos):
        self.document.add_heading('Top ventas', 2)
        self.document.add_paragraph(
            'Producto más vendido: '+producto_mas_vendido, style='List Bullet'
        )
        self.document.add_paragraph(
            'Producto con más ingresos: '+producto_mas_ingresos, style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag7(self,campañas , cuadro_texto, graficos):
        self.document.add_heading('Advertising', 1)
        self.document.add_heading('Campañas', 2)
        for c in campañas:
            self.document.add_paragraph(
                c, style='List Bullet'
            )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag8(self, ingreso_semanal, ingreso_anterior_semana,  gasto_semanal, gasto_anterior_semana, cuadro_texto, graficos):
        self.document.add_heading('Ingresos semanales', 2)
        self.document.add_paragraph(
            'Ingreso semanal: '+str(ingreso_semanal)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Ingreso semana anterior: '+str(ingreso_anterior_semana)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Gasto semanal: '+str(gasto_semanal)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Gasto semana anterior: '+str(gasto_anterior_semana)+' €', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    

    def finaliza_doc(self):
        self.document.add_page_break()
        self.document.save(os.path.dirname(__file__)+'/demo.docx')

informe_sem=informe_semanal('picsil')
informe_sem.portada()
informe_sem.pag1(4452.5, 20, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
informe_sem.pag2(16284.87, 100, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
informe_sem.pag3(405, 110.25, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
informe_sem.pag4('Esto es un texto para editar',['ingresos.png', 'ventas.png'], exceso_inventario=10, reabastecer=20)
informe_sem.pag5(98, 95.2, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
informe_sem.pag6('B07IEURE', 'B06DFE45', 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
informe_sem.pag7(['Campaña 1', 'Campaña 2', 'Campaña 3', 'Campaña 4'], 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
informe_sem.pag8(1000, 950, 100, 99, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])

informe_sem.finaliza_doc()
