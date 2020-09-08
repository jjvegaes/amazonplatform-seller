from docx import Document
from docx.shared import Inches
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
#pip install python-docx

#clase que crea el documento word
class informe_semanal():
    #Constructor que se le pasa el nombre de la empresa
    def __init__(self, empresa):
        self.empresa=empresa
        self.document = Document()

    #Crea la portada del documento
    def portada(self):
        self.document.add_paragraph('\n\n\n\n\n\n\n')#Saltos de línea
        t=self.document.add_heading('Informe Semanal '+self.empresa, 0)#Título
        
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER#Texto centrado
        #Tamaño del texto:
        title_style = t.style
        title_style.font.size = Pt(40)

        self.document.add_paragraph('\n\n')
        #Logo Macreif:
        self.document.add_picture(os.path.dirname(__file__)+'/imagenes/logo_mcreif.png', width=Inches(4.5))
        last_paragraph = self.document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()#Pasa a otra página

    #A partir de ahora se definen las distintas páginas del informe, tienen todas una estructura parecida

    def pag1(self, ingresos, ventas, cuadro_texto, graficos):
        self.document.add_heading('Temas generales', 1)#Título nivel 1
        self.document.add_heading('Ingresos semanales', 2)#Título nivel 2
        #Lista con los ingresos y ventas:
        self.document.add_paragraph(
            'Ingresos: '+str(ingresos)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Ventas: '+str(ventas)+' unidades', style='List Bullet'
        )
        #Parafo para añadir cosas:
        self.document.add_paragraph(cuadro_texto)
        #Aquí se insertan los gráficos:
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag2(self, ingresos, ventas, cuadro_texto, graficos):
        self.document.add_heading('Ingresos mensuales', 2)
        self.document.add_paragraph(
            'Ingresos: '+str(ingresos)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Ventas: '+str(ventas)+' unidades', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag2_2(self, ticket_medio, cuadro_texto, graficos):
        self.document.add_heading('Ticket Medio', 2)
        self.document.add_paragraph(
            'Ticket medio de compra: '+str(ticket_medio)+' €', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag3(self, obj_sell_in, obj_sell_out, cuadro_texto, graficos):
        self.document.add_heading('Objetivo sell in y sell out', 2)
        self.document.add_paragraph(
            'Objetivo sell in: '+str(obj_sell_in)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Objetivo sell out: '+str(obj_sell_out)+' €', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag3_2(self, unidades_totales, unidades_aptas, cuadro_texto, graficos):
        self.document.add_heading('Inventario', 1)
        self.document.add_heading('Unidades aptas en el inventario', 2)
        self.document.add_paragraph(
            'Unidades totales en el inventario: '+str(unidades_totales)+' unidades', style='List Bullet'
        )
        self.document.add_paragraph(
            'Unidades aptas para la venta del inventario: '+str(unidades_aptas)+' unidades', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag4(self, cuadro_texto, graficos, buy_box=None, exceso_inventario=None, numero_low_stock=None, numero_out_of_stock=None):
        self.document.add_heading('Alertas', 2)
        if buy_box != None:
            self.document.add_paragraph(
                'Buy box perdida: '+str(buy_box)+' %', style='List Bullet'
            )
        if exceso_inventario != None:
            self.document.add_paragraph(
                'Unidades en exceso del inventario: '+str(exceso_inventario)+' unidades', style='List Bullet'
            )
        if numero_low_stock != None:
            self.document.add_paragraph(
                'Productos que tienen bajo stock: '+str(numero_low_stock)+' unidades', style='List Bullet'
            )
        if numero_out_of_stock != None:
            self.document.add_paragraph(
                'Productos que no tienen stock: '+str(numero_out_of_stock)+' unidades', style='List Bullet'
            )
        
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag5(self, buy_box_stock, buy_box_precio, cuadro_texto, graficos):
        self.document.add_heading('Posicionamiento', 1)
        self.document.add_heading('Buy box', 2)
        self.document.add_paragraph(
            'Buy box perdida por stock : '+str(buy_box_stock)+' %', style='List Bullet'
        )
        self.document.add_paragraph(
            'Buy box perdida por precio: '+str(buy_box_precio)+' %', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag6(self, producto_mas_ingresos, producto_mas_vendido, cuadro_texto, graficos):
        self.document.add_heading('Top ventas', 2)
        self.document.add_paragraph(
            'Producto más vendido: '+producto_mas_vendido, style='List Bullet'
        )
        self.document.add_paragraph(
            'Producto con más ingresos: '+producto_mas_ingresos, style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(6.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()
    
    def pag6_2(self, producto_mas_ingresos, producto_mas_vendido, cuadro_texto, graficos):
        self.document.add_heading('Top ventas mensuales', 2)
        self.document.add_paragraph(
            'Producto más vendido: '+producto_mas_vendido, style='List Bullet'
        )
        self.document.add_paragraph(
            'Producto con más ingresos: '+producto_mas_ingresos, style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(6.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag7(self,campañas , cuadro_texto, graficos):
        self.document.add_heading('Advertising', 1)
        self.document.add_heading('Campañas', 2)
        for c in campañas:
            self.document.add_paragraph(
                c, style='List Bullet'
            )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag8(self, ingreso_semanal, ingreso_anterior_semana,  gasto_semanal, gasto_anterior_semana, cuadro_texto, graficos):
        self.document.add_heading('Ingresos semanales', 2)
        self.document.add_paragraph(
            'Ingreso semanal: '+str(ingreso_semanal)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Ingreso semana anterior: '+str(ingreso_anterior_semana)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Gasto semanal: '+str(gasto_semanal)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Gasto semana anterior: '+str(gasto_anterior_semana)+' €', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag9(self, ingreso_mensual,  gasto_mensual,  cuadro_texto, graficos):
        self.document.add_heading('Ingresos mensuales', 2)
        self.document.add_paragraph(
            'Ingreso mensual: '+str(ingreso_mensual)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Gasto mensual: '+str(gasto_mensual)+' €', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag10(self, ingreso_anual,  gasto_anual,  cuadro_texto, graficos):
        self.document.add_heading('Ingresos anuales', 2)
        self.document.add_paragraph(
            'Ingreso mensual: '+str(ingreso_anual)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'Gasto mensual: '+str(gasto_anual)+' €', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag11(self, roas,  acos,  cuadro_texto, graficos):
        self.document.add_heading('ROAS y ACOS', 2)
        self.document.add_paragraph(
            'ROAS: '+str(roas)+' €', style='List Bullet'
        )
        self.document.add_paragraph(
            'ACOS: '+str(acos)+' %', style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def pag12(self, pedidos_totales,  pedidos_pendientes, pedidos_enviando, pedidos_enviados, pedidos_cancelados,  cuadro_texto, graficos):
        self.document.add_heading('Pedidos', 1)
        self.document.add_paragraph(
            'Pedidos totales: '+str(pedidos_totales), style='List Bullet'
        )
        self.document.add_paragraph(
            'Pedidos pendientes: '+str(pedidos_pendientes), style='List Bullet'
        )
        self.document.add_paragraph(
            'Pedidos enviando: '+str(pedidos_enviando), style='List Bullet'
        )
        self.document.add_paragraph(
            'Pedidos enviados: '+str(pedidos_enviados), style='List Bullet'
        )
        self.document.add_paragraph(
            'Pedidos cancelados: '+str(pedidos_cancelados), style='List Bullet'
        )
        self.document.add_paragraph(cuadro_texto)
        for gr in graficos:
            self.document.add_picture(os.path.dirname(__file__)+'/imagenes/'+gr, width=Inches(4.5))
            last_paragraph = self.document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    

    def finaliza_doc(self):
        self.document.add_page_break()
        self.document.save(os.path.dirname(__file__)+'/demo.docx')

#informe_sem=informe_semanal('picsil')
#informe_sem.portada()
#informe_sem.pag1(4452.5, 20, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
#informe_sem.pag2(16284.87, 100, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
#informe_sem.pag3(405, 110.25, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
#informe_sem.pag4('Esto es un texto para editar',['ingresos.png', 'ventas.png'], exceso_inventario=10, reabastecer=20)
#informe_sem.pag5(98, 95.2, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
#informe_sem.pag6('B07IEURE', 'B06DFE45', 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
#informe_sem.pag7(['Campaña 1', 'Campaña 2', 'Campaña 3', 'Campaña 4'], 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
#informe_sem.pag8(1000, 950, 100, 99, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
#informe_sem.pag9(1800, 1000, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
#informe_sem.pag10(54645, 9504, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
#informe_sem.pag11(126, 15, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
#informe_sem.pag12(1000, 950, 50, 'Esto es un texto para editar', ['ingresos.png', 'ventas.png'])
#informe_sem.finaliza_doc()
